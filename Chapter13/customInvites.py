#!python3 
#customInvites - write a unique invite for each guest from guest.txt file

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
logging.basicConfig(level=logging.DEBUG, format=' %(asctime)s - %(levelname)s- %(message)s')
#logging.disable(logging.CRITICAL)

def main():
	textFile = open('guests.txt','r')
	docFile = docx.Document()
	#removes the whitespace('including \n') from the text and adds the guest in the guestList
	guestList = [text.rstrip() for text in textFile]
	for guest in guestList:
		paragraph = docFile.add_paragraph('It would be a pleasure to have the company of')
		paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		paragraph.runs[0].font.name = 'Brush Script MT'
		paragraph.runs[0].font.size = docx.shared.Pt(24)
		paragraph.paragraph_format.space_after =docx.shared.Pt(0)
		
		paragraph = docFile.add_paragraph(str(guest))
		paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		paragraph.runs[0].font.name = 'Comic Sans MS'
		paragraph.runs[0].font.size = docx.shared.Pt(20)
		paragraph.paragraph_format.space_after =docx.shared.Pt(0)	

		paragraph = docFile.add_paragraph('At 11010 Memory Lane on the Evening of')
		paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		paragraph.runs[0].font.name = 'Brush Script MT'
		paragraph.runs[0].font.size = docx.shared.Pt(24)
		paragraph.paragraph_format.space_after =docx.shared.Pt(0)
		
		paragraph = docFile.add_paragraph('April 1')
		paragraph.add_run('st')
		paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		for run in paragraph.runs:
			run.font.name = 'Comic Sans MS'
			run.font.size = docx.shared.Pt(20)	
		paragraph.runs[1].font.superscript = True
		paragraph.paragraph_format.space_after =docx.shared.Pt(0)
		
		
		paragraph = docFile.add_paragraph('At 7 o\'clock')
		paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		paragraph.runs[0].font.name = 'Brush Script MT'
		paragraph.runs[0].font.size = docx.shared.Pt(24)
		paragraph.paragraph_format.space_after =docx.shared.Pt(0)
		docFile.add_page_break()
	
	#remove the last break page.
	docFile.paragraphs[len(docFile.paragraphs)-1].text=None
	docFile.save('guestInvites.docx')	
	logging.debug('Complete')
	
if __name__=='__main__':
	main()